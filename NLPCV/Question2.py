"""
CDAC NLP & Computer Vision -- Combined Script
Combined from all 32 notebooks in mituskillologies/cdac-nlpcv-may26
https://github.com/mituskillologies/cdac-nlpcv-may26

Structure:
  1. A consolidated import block below (every package used anywhere in the course).
  2. Each original notebook follows as its own section (## headers), with its
     own original imports kept in place too.

Notes:
  - Lines that were Jupyter shell/magic commands (e.g. 'pip install ...',
    '!wget ...') are commented out with a [shell/magic] marker -- run those
    manually in a terminal or notebook cell.
  - A few cells used notebook-only constructs and are fully commented out with
    a NOTE marker; open the original notebook (or the combined .ipynb) for those.
  - Some sections need extra setup before they will run (nltk.download(...),
    spaCy model downloads, local files like PDFs/Word docs/GloVe vectors).
  - Run sections independently rather than the whole file top-to-bottom, since
    several topics reuse the same variable names (e.g. text, model, df).
"""

# ===================== CONSOLIDATED IMPORTS =====================
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
import string
from nltk.tokenize import WhitespaceTokenizer
from nltk.tokenize import SpaceTokenizer
from nltk.tokenize import TabTokenizer
from nltk.tokenize import LineTokenizer
from nltk.tokenize import MWETokenizer
from nltk.tokenize import TweetTokenizer
import re
from nltk.tokenize import word_tokenize
from nltk import pos_tag
from nltk.corpus import indian
from nltk import TnT
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer
from nltk.stem import LancasterStemmer
from nltk.stem import WordNetLemmatizer
import spacy
from spacy import displacy
from nltk import pos_tag, RegexpParser
from nltk import ne_chunk
from nltk.corpus import wordnet
import pyiwn
from nltk.wsd import lesk
import language_tool_python
import requests
from bs4 import BeautifulSoup
from nltk import FreqDist
from pypdf import PdfReader
import docx
import os
from sumy.summarizers.text_rank import TextRankSummarizer
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer
from nltk.tokenize import sent_tokenize
from sumy.summarizers.lex_rank import LexRankSummarizer
from sumy.summarizers.lsa import LsaSummarizer
from transformers import pipeline
import transformers
import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.feature_extraction.text import TfidfVectorizer
import seaborn as sns
from nltk.stem import LancasterStemmer, WordNetLemmatizer, PorterStemmer
from sklearn.model_selection import train_test_split
from sklearn.svm import SVC
from sklearn.metrics import accuracy_score, classification_report
from sklearn.metrics import ConfusionMatrixDisplay
import joblib
import re, string
from wordcloud import WordCloud
from sklearn.neighbors import NearestCentroid
from sklearn.metrics import classification_report, accuracy_score
from sklearn.cluster import KMeans
from kneed import KneeLocator
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
import nltk, re, string, os
from keras.models import Sequential
from keras.layers import Input, Dense
from keras.utils import to_categorical, plot_model
from keras.layers import Input, Dense, Conv1D, MaxPool1D, Embedding, Flatten
from keras.utils import to_categorical
from keras.layers import SimpleRNN, Input, Dense
import tensorflow_datasets as tfds
import tensorflow as tf
from keras.layers import Dense, Embedding, SimpleRNN, Bidirectional, LSTM, GRU
from keras.layers import Dense, Embedding, LSTM, GRU
import wave, math, struct
from keras.layers import Dense, LSTM, GRU, Activation, Input
import gensim
from gensim.models import Word2Vec
from sklearn.decomposition import PCA
from sklearn.metrics.pairwise import cosine_similarity
from gensim.models import KeyedVectors
from gensim.scripts.glove2word2vec import glove2word2vec
from transformers import Conversation


# ======================================================================
# 01 Tokenization
# ======================================================================
# #### install the libraries

# [shell/magic - run manually]  pip install nltk spacy textblob -U

import nltk

nltk.download()

nltk.download('punkt')
nltk.download('punkt_tab')  # tokenization
nltk.download('stopwords')  # stopwords
nltk.download('averaged_perceptron_tagger')  # POS Tagging
nltk.download('averaged_perceptron_tagger_eng')  # POS Tagging
nltk.download('punkt')
nltk.download('wordnet')  # wordnet and lemmatization
nltk.download('indian')  # INDIAN POS Tagging
nltk.download('omw-1.4')  # multilinguial wordnet

# #### Sample Task

sent = 'They told that their ages are 24, 31 and 29 respectively.'
# Find the average of ages mentioned in this sentence

words = sent.split()
ages = []
for word in words:
    if word.isdigit(): 
        ages.append(int(word))

ages

sum(ages) / len(ages)

ages = [int(word) for word in sent.split() if word.isdigit()]

ages

# #### Word and sentence tokenizer

# import the tokenizer function
from nltk.tokenize import word_tokenize, sent_tokenize

sent = '''Hello friends! How are you?
Welcome to the world of Python Programming.'''

print(sent)

sents = sent_tokenize(sent)

sents

sents[0]

sents[1]

sents[2]

words = word_tokenize(sent)

words

# Among all tokens, find percentage of punctuation symbols in it

import string

string.punctuation

count_punct = 0
for word in word_tokenize(sent):
    if word in string.punctuation: 
        count_punct += 1

count_punct

count_punct / len(word_tokenize(sent)) * 100

sent

new_str = ''
for char in sent: 
    if char not in string.punctuation: 
        new_str += char    

print(new_str)

str1 = '''पुणे हे भारताच्या दख्खन पठारावरील महाराष्ट्राचे एक महत्त्वाचे शहर आहे. पुणे जिल्ह्याचे आणि पुणे विभागाचे हे प्रशासकीय मुख्यालय आहे. २०११ जनगणनेनुसार, शहराच्या हद्दीतील ३.१ दशलक्ष लोकसंख्येसह पुणे हे भारतातील नववे सर्वात जास्त लोकसंख्या असलेले शहर आहे. महानगर प्रदेशातील ७.२ दशलक्ष रहिवासी लोकसंख्या आहे, यानुसार ते आठवे सर्वाधिक लोकसंख्या असलेले शहर आहे.[८] पुणे शहर पुणे महानगर प्रदेशाचा एक भाग आहे.[९] भारतातील सर्वात मोठ्या आयटी हबपैकी पुणे एक आहे.[१०][११] हे भारतातील सर्वात महत्वाचे ऑटोमोबाईल आणि उत्पादन केंद्रांपैकी एक आहे. पुणे हे एक विकसनशील शहर आहे.[१२] या शहराला येथील उच्च प्रतिष्ठित शैक्षणिक संस्थांमुळे "पूर्वेकडील ऑक्सफर्ड" म्हणून संबोधले जाते.[१३][१४][१५] "भारतातील सर्वात राहण्यायोग्य शहर" म्हणून अनेक वेळा पुण्याला स्थान देण्यात आले आहे. [१६][१७]
पुण्यावर राष्ट्रकूट राजघराणे, अहमदनगर सल्तनत, मुघल, आदिल शाही घराणे यांनी राज्य केले आहे. १८ व्या शतकात हे शहर मराठा साम्राज्याचा भाग होते, आणि मराठा साम्राज्याचे पंतप्रधान, पेशव्यांची हे आसन होते.[१८] पाताळेश्वर लेणी, शनिवारवाडा, शिंदे छत्री, विश्रामबाग वाडा यासारख्या अनेक ऐतिहासिक खुणा या कालखंडातील आहेत. वेगवेगळ्या कालखंडातील ऐतिहासिक स्थळे शहरभर पसरलेली आहेत.
'''

print(str1)

word_tokenize(str1)

new_sent = 'Dr. Mrs. Kulkarni got 12 lakh rupees salary.'

word_tokenize(new_sent)

# #### File operations

# Load the file mydata.txt and find average of numbers

f = open('mydata.txt')
data = f.read()

data = word_tokenize(data)

nums = []
for word in data:
    if not word.isalpha() and word not in string.punctuation: 
        nums.append(float(word))

sum(nums) / len(nums)

nums

# #### White Space Tokenizer

from nltk.tokenize import WhitespaceTokenizer

tk = WhitespaceTokenizer()

sent

tk.tokenize(sent)

# #### Space Tokenizer

from nltk.tokenize import SpaceTokenizer

tk = SpaceTokenizer()

tk.tokenize(sent)

sent = '''Hello friends! How are you?
Welcome\tto the world of Python\tProgramming.'''

print(sent)

# #### Tab Tokenizer

from nltk.tokenize import TabTokenizer

tk = TabTokenizer()

tk.tokenize(sent)

sent.split('\t')

# #### Line Tokenizer

from nltk.tokenize import LineTokenizer

tk = LineTokenizer()

tk.tokenize(sent)

sent.split('\n')

# #### Multi Word Expression Tokenizer

from nltk.tokenize import MWETokenizer

nsent = 'Chhatrapati Shivaji Maharaj was founder of Maratha empire'

word_tokenize(nsent)

mwe = MWETokenizer(separator=' ')

mwe.add_mwe(('Chhatrapati','Shivaji'))

mwe.tokenize(word_tokenize(nsent))

# #### Tweet Tokenizer

name = 'तुषार कुटे'

name

name.replace('ष','श')

name.startswith('त')

name.split()

name.isascii()

name.index('र')

char = '\u0915\u093E'
char

char = '\u0041\u0073'
char

ord('क')

chr(2354)

sent = 'Hello friends!:) How are you? Welcome to Python <3Programming.'

sent

from nltk.tokenize import TweetTokenizer

tk = TweetTokenizer()

tk.tokenize(sent)

text = '😀'

print(text)

data = 'Hello friends👋 How are you?💥 Welcome to 🌎Python Programming.💻'

print(data)

word_tokenize(data)

tk.tokenize(data)

# #### Custom Tokenizer

import re

def custom_tokenizer(text):
    return re.split(r"[.,;?!\s]+", text)

text = "This is some text with punctuation > Let's tokenize it. Is it ok?"

custom_tokenizer(text)

class CommaTokenizer: 
    def tokeinze(self, text):
        return re.split(r"[,\s]+", text)

tk = CommaTokenizer()

fruits = 'apple, banana, cherry, pineapple.'

tk.tokeinze(fruits)




# ======================================================================
# 02 POS Tagging
# ======================================================================
from nltk.tokenize import word_tokenize
from nltk import pos_tag

sent = 'A quick brown fox jumps over the lazy dog.'

tags = pos_tag(word_tokenize(sent))

tags

data = 'Pune (Marathi: Puṇē, pronounced [ˈpuɳe] ⓘ), previously spelled in English as Poona (the official name until 1978),[16][17] is a city in the state of Maharashtra in the Deccan Plateau in Western India. It is the administrative headquarters of the Pune district, and of Pune division. In terms of the total amount of land under its jurisdiction, Pune is the largest city in Maharashtra by area, with a geographical area of 516.18km2,[18] though by population it comes in a distant second to Mumbai. According to the 2011 Census of India, Pune has 7.2 million residents in the metropolitan region, making it the seventh-most populous metropolitan area in India.[19] The city of Pune is part of Pune Metropolitan Region.[20] Pune is one of the largest IT hubs in India.[21][22] It is also one of the most important automobile and manufacturing hubs of India. Pune is often referred to as the "Oxford of the East" because of its educational institutions.[23][24][25] It has been ranked "the most liveable city in India" several times.'

clean = [word for word in word_tokenize(data) 
         if word.isalpha() or word.isdigit()]

tags = pos_tag(clean)

tags

for word, tag in tags:
    if tag.startswith('JJ'): 
        print(word)

# Find and print all the pairs of adjective + noun 

pairs = []
for i in range(len(tags)):
    if tags[i][1].startswith('JJ') and tags[i+1][1].startswith('NN'):
        pairs.append([tags[i][0], tags[i+1][0]])

pairs




# ======================================================================
# 03 POS Tagging - Indian Languages
# ======================================================================
from nltk.tokenize import word_tokenize
import nltk

nltk.download('indian')

from nltk.corpus import indian

indian.fileids()

indian.sents()

for file in indian.fileids():
    print(file)
    print(len(indian.words(file)))

indian.sents('marathi.pos')

indian.tagged_sents('marathi.pos')

# Import the tagger class
from nltk import TnT

# create the object
tagger = TnT()

# extract the tags
tags = indian.tagged_sents('marathi.pos')

# train the tagger 
tagger.train(tags)

new_sent = 'हा सामना सोमवारी उशिरा सुरू झाला.'

tagger.tag(word_tokenize(new_sent))




# ======================================================================
# 04 Stopwords Removal
# ======================================================================
import nltk

nltk.download('stopwords')

from nltk.corpus import stopwords

swords = stopwords.words('english')

swords

sent = 'Hello friends! How are you? Welcome to world of Python Programming.'

# List of words without stopwords

from nltk.tokenize import word_tokenize
import string

clean_words = []
for word in word_tokenize(sent):
    if word.lower() not in swords and word not in string.punctuation: 
        clean_words.append(word)

clean_words




# ======================================================================
# 06 Stemming
# ======================================================================
words1 = 'working','work','worked','workable'
words2 = 'bigger', 'biggest', 'big'
words3 = 'consulting','consultation','consulted','consultative'
words4 = 'player', 'played', 'players', 'playing', 'plays'

from nltk.stem import PorterStemmer

ps = PorterStemmer()

for word in words1:
    print(word, '->', ps.stem(word))

for word in words3:
    print(word, '->', ps.stem(word))

for word in words2:
    print(word, '->', ps.stem(word))

for word in words4:
    print(word, '->', ps.stem(word))

from nltk.stem import LancasterStemmer

ls = LancasterStemmer()

for word in words1:
    print(word, '->', ls.stem(word))

for word in words2:
    print(word, '->', ls.stem(word))

for word in words3:
    print(word, '->', ls.stem(word))

for word in words4:
    print(word, '->', ls.stem(word))

sent = '''Hello friends! How are you? My friend like to program 
in Python Programming language.'''

# remove stopwords, punctuations and suffixes and print the list of words

from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
import string

swords = stopwords.words('english')

clean_words = []
for word in word_tokenize(sent):
    if word.lower() not in swords and word not in string.punctuation: 
        clean_words.append(ps.stem(word))

clean_words




# ======================================================================
# 07 Lemmatization
# ======================================================================
from nltk.stem import WordNetLemmatizer

wnl = WordNetLemmatizer()

words1 = 'running', 'ran', 'runs', 'runnable', 'run'
words2 = 'consulting', 'consulted', 'consultation', 'consult'
words3 = 'big', 'biggest', 'bigger'
words4 = 'players', 'plays', 'playing', 'played'
words5 = 'eaten', 'ate', 'eat', 'eating'

for word in words1: 
    print(word,"\t->", wnl.lemmatize(word, pos='v'))

for word in words2: 
    print(word,"\t->", wnl.lemmatize(word, pos='v'))

for word in words3: 
    print(word,"\t->", wnl.lemmatize(word, pos='a'))

for word in words4: 
    print(word,"\t->", wnl.lemmatize(word, pos='v'))

for word in words4: 
    print(word,"\t->", wnl.lemmatize(word, pos='n'))

for word in words5: 
    print(word,"\t->", wnl.lemmatize(word, pos='v'))

data = 'Pune (Marathi: Puṇē, pronounced [ˈpuɳe] ⓘ), previously spelled in English as Poona (the official name until 1978),[16][17] is a city in the state of Maharashtra in the Deccan Plateau in Western India. It is the administrative headquarters of the Pune district, and of Pune division. In terms of the total amount of land under its jurisdiction, Pune is the largest city in Maharashtra by area, with a geographical area of 516.18km2,[18] though by population it comes in a distant second to Mumbai. According to the 2011 Census of India, Pune has 7.2 million residents in the metropolitan region, making it the seventh-most populous metropolitan area in India.[19] The city of Pune is part of Pune Metropolitan Region.[20] Pune is one of the largest IT hubs in India.[21][22] It is also one of the most important automobile and manufacturing hubs of India. Pune is often referred to as the "Oxford of the East" because of its educational institutions.[23][24][25] It has been ranked "the most liveable city in India" several times.'

# extract all distinct root forms of verbs from this.

from nltk.tokenize import word_tokenize
from nltk import pos_tag
import string

verbs = set()
tags = pos_tag(word_tokenize(data))
for word, tag in tags:
    if word.isalpha() and tag.startswith('V'):
        verbs.add(wnl.lemmatize(word.lower(), pos = 'v'))

verbs




# ======================================================================
# 08 Named Entity Recognition
# ======================================================================
import spacy

# [shell/magic - run manually]  !python3 -m spacy download en_core_web_sm

nlp = spacy.load('en_core_web_sm')

nlp

sent = nlp('''Mark Zukerberg will meet Aditya Joshi
on Monday 25 November 2025, 10am for $3 Trillion deal
at Pune for second time.''')

sent

sent.ents

for ent in sent.ents:
    print(ent.text,' ---> ', ent.label_)

spacy.explain('PERSON')

spacy.explain('DATE')

spacy.explain('MONEY')

spacy.explain('ORDINAL')

mytext = '''Sachin Ramesh Tendulkar (/ˌsʌtʃɪn tɛnˈduːlkər/ ⓘ; Marathi: [sətɕin t̪eɳɖulkəɾ]; born 24 April 1973) is an Indian former international cricketer who captained the Indian national team. Often dubbed the "God of Cricket" in India, he is widely regarded as one of the greatest cricketers of all time.[5] He holds several world records, including being the all-time highest run-scorer in international cricket,[6] receiving the most player of the match awards in international cricket,[7] and being the only batsman to score 100 international centuries.[8] Tendulkar was a Member of Parliament, Rajya Sabha by presidential nomination from 2012 to 2018.[9][10]
Tendulkar took up cricket at the age of eleven, made his Test match debut on 15 November 1989 against Pakistan in Karachi at the age of sixteen, and went on to represent Mumbai domestically and India internationally for over 24 years.[11] In 2002, halfway through his career, Wisden ranked him the second-greatest Test batsman of all time, behind Don Bradman, and the second-greatest ODI batsman of all time, behind Viv Richards.[12] The same year, Tendulkar was a part of the team that was one of the joint-winners of the 2002 ICC Champions Trophy. Later in his career, Tendulkar was part of the Indian team that won the 2011 Cricket World Cup, his first win in six World Cup appearances for India.[13] He had previously been named "Player of the Tournament" at the 2003 World Cup.'''

sent = nlp(mytext)

for ent in sent.ents:
    print(ent.text,' ---> ', ent.label_)

from spacy import displacy

displacy.render(sent, style='ent', jupyter=True)




# ======================================================================
# 09 Chunking
# ======================================================================
from nltk.tokenize import word_tokenize
from nltk import pos_tag, RegexpParser

# Create the chunk grammer for regular expression
grammer = "NP:{<DT><JJ><NN>}"

sent = 'It was a beautiful game by him. An amazing play was witnessed.'

parser = RegexpParser(grammer)

tree = parser.parse(pos_tag(word_tokenize(sent)))

tree

tree.draw()

grammer = r"NP:{<DT>?<JJ>*<NN>+}"
# <DT>? Zero or One Determiner 
# <JJ>* Zero or more adjectives 
# <NN>+ One or nouns

sent = 'It was a beautiful fine game by him. An amazing play was witnessed.'
parser = RegexpParser(grammer)
tree = parser.parse(pos_tag(word_tokenize(sent)))

tree

sent = '''It was a beautiful fine game by him. 
An amazing play was witnessed. They won the game.'''
parser = RegexpParser(grammer)
tree = parser.parse(pos_tag(word_tokenize(sent)))
tree

# #### NER Chunking

from nltk import ne_chunk

sent = 'Mark Zukerberg is CEO of Meta at California.'

tree = ne_chunk(pos_tag(word_tokenize(sent)))

tree

tree.leaves()




# ======================================================================
# 10 Wordnet
# ======================================================================
import nltk
nltk.download('wordnet')

from nltk.corpus import wordnet

wordnet.synsets('create')

synonyms = set()

for word in wordnet.synsets('create'):
    synonyms.add(word.lemmas()[0].name())

synonyms

synonyms = set()

for word in wordnet.synsets('association'):
    synonyms.add(word.lemmas()[0].name())

synonyms

# #### Definitions and Examples

word = wordnet.synsets('post')
print("Word and POS: ", word[0].name())
print("Synonyms:", word[0].lemmas()[1].name())
print("Meaning:", word[0].definition())
print("Examples:", word[0].examples())

word = wordnet.synsets('post')
print("Word and POS: ", word[1].name())
print("Synonyms:", word[1].lemmas()[1].name())
print("Meaning:", word[1].definition())
print("Examples:", word[1].examples())

word = wordnet.synsets('post')
print("Word and POS: ", word[2].name())
print("Synonyms:", word[2].lemmas()[1].name())
print("Meaning:", word[2].definition())
print("Examples:", word[2].examples())

for word in wordnet.synsets('sound'): 
    print(word.name(),":",word.definition())

for word in wordnet.synsets('bat'): 
    print(word.name(),":",word.definition())

for word in wordnet.synsets('bat', pos='n'): 
    print(word.name(),":",word.definition())

# Extract all nouns, verbs and adjective synonyms for a word

nouns = set()
verbs = set()
adjs = set()
for word in wordnet.synsets('sound', pos='n'):
    nouns.add(word.lemmas()[0].name())
for word in wordnet.synsets('sound', pos='v'):
    verbs.add(word.lemmas()[0].name())
for word in wordnet.synsets('sound', pos='a'):
    adjs.add(word.lemmas()[0].name())

nouns

verbs

adjs

# #### Extract the antonyms

antonyms = set()

for word in wordnet.synsets('possible'):
    for lemma in word.lemmas():
        if lemma.antonyms():
            antonyms.add(lemma.antonyms()[0].name())

antonyms

antonyms = set()

for word in wordnet.synsets('computer'):
    for lemma in word.lemmas():
        if lemma.antonyms():
            antonyms.add(lemma.antonyms()[0].name())

antonyms

# #### Extract the hypernyms

word = wordnet.synsets('dog')
for x in word:
    hypernyms = x.hypernyms()
    for h in hypernyms:
        print(h.name())

# #### Extract the meronyms

word = wordnet.synsets('computer')
for x in word:
    meronyms = x.part_meronyms()
    for m in meronyms:
        print(m.name())

# Extract the holonyms

word = wordnet.synset('leaf.n.03')

print(word.part_holonyms())




# ======================================================================
# 11 Indo Wordnet
# ======================================================================
# [shell/magic - run manually]  pip install pyiwn

import pyiwn

iwn = pyiwn.IndoWordNet(lang=pyiwn.Language.MARATHI)

word = iwn.synsets('धरती')

word

# Explore the specific synsets

word[0].synset_id()

word[0].pos()

# meaning
word[0].gloss()

# Example in senetence
word[0].examples()

# synonyms
word[0].lemma_names()

# Hyponyms

hyponyms = iwn.synset_relation(word[0], pyiwn.SynsetRelations.HYPONYMY)

for h in hyponyms: 
    print(h.lemma_names()[0])

word = iwn.synsets('आंबा')
hyponyms = iwn.synset_relation(word[0], pyiwn.SynsetRelations.HYPONYMY)
for h in hyponyms: 
    print(h.lemma_names()[0])




# ======================================================================
# 12 Word Sense Disambiguition
# ======================================================================
import nltk
from nltk.tokenize import word_tokenize
from nltk.wsd import lesk

sent1 = 'I went to bank to deposit my money.'
sent2 = 'We pitched our tent by the bank of river.'

sense1 = lesk(word_tokenize(sent1), 'bank', 'n')

sense2 = lesk(word_tokenize(sent2), 'bank', 'n')

sense1.definition()

sense2.definition()

nltk.__version__

sent3 = 'I spread sweet strawberry jam on my toast for breakfast.'
sent4 = 'I was late to work because I was stuck in massive traffic jam.'

sense3 = lesk(word_tokenize(sent3), 'jam', 'n')
sense4 = lesk(word_tokenize(sent4), 'jam', 'n')

sense3.definition()

sense4.definition()

sent5 = 'The small bat flew out of the dark cave and into the night sky.'
sent6 = 'The cricket player swung the wooden bat and hit a massive six.'

sense5 = lesk(word_tokenize(sent5), 'bat', 'n')
sense6 = lesk(word_tokenize(sent6), 'bat', 'n')

sense5.definition()

sense6.definition()






# ======================================================================
# 13 Syntactic Parsing, Grammer Checking
# ======================================================================
# [shell/magic - run manually]  pip install language-tool-python

import language_tool_python

# download the support for language

mytool = language_tool_python.LanguageTool('en-US')

# define the sentence
my_text = 'He have collected a documents. When I is there.'

matches = mytool.check(my_text)

len(matches)

for m in matches:
    print(m)

mytool.correct(my_text)




# ======================================================================
# 14 Information Extraction - Web Scraping
# ======================================================================
import requests

url = 'https://en.wikipedia.org/wiki/Rajgad'

response = requests.get(url)

response

headers = {
    'User-Agent':'Custom-Script/1.9 (tushar@gmail.com)', 
    'Accept-Language': 'en-US,en;q=0.5'
}

response = requests.get(url, headers=headers)

response

print(response.text)

# [shell/magic - run manually]  pip install bs4

from bs4 import BeautifulSoup

soup = BeautifulSoup(response.text, 'html.parser')

type(soup)

soup.find('title')

text = soup.get_text()

print(text)

with open('rajgad.txt','w') as f:
    f.write(text)

# #### Find top 10 words occured frequently in this article

from nltk.tokenize import word_tokenize

tokens = word_tokenize(text)

from nltk import FreqDist

freq = FreqDist(tokens)

freq.plot(10)

from nltk.corpus import stopwords
from nltk.stem import LancasterStemmer
ls = LancasterStemmer()

tokens1 = [ls.stem(token.lower()) for token in tokens 
           if token.isalpha() and token.lower() 
           not in stopwords.words('english')]

freq = FreqDist(tokens1)
freq.plot(20)

freq.most_common(20)

swords = stopwords.words('english') + ['retrieved','archived','edit','b','original']

url = 'https://en.wikipedia.org/wiki/Saturn'
response = requests.get(url, headers=headers)
soup = BeautifulSoup(response.text, 'html.parser')
text = soup.get_text()
tokens = word_tokenize(text)
tokens1 = [token.lower() for token in tokens 
           if token.isalpha() and token.lower() 
           not in swords]
freq = FreqDist(tokens1)
freq.plot(20)






# ======================================================================
# 15 Informaton Extraction - Documents
# ======================================================================
# #### Read from the pdf

# [shell/magic - run manually]  pip install pypdf

# #### Source files

# https://mitu.co.in/dataset 
# Download-> course.pdf, spp.docx and store in current working directory

from pypdf import PdfReader

reader = PdfReader('datasets/course.pdf')

reader

len(reader.pages)

reader.pages[0]

page1 = reader.pages[0].extract_text()

print(page1)

pgno = 1
for page in reader.pages:
    print("Page No.", pgno)
    print(page.extract_text())
    pgno += 1

text =''
for page in reader.pages:
    text += page.extract_text()

with open('course.txt','w') as f:
    f.write(text)

# #### Reading data from word documents

# [shell/magic - run manually]  pip install python-docx

import docx

doc = docx.Document('datasets/spp.docx')

doc

len(doc.paragraphs)

for para in doc.paragraphs: 
    print(para.text)

# Extract the bold text 
for para in doc.paragraphs: 
    for run in para.runs:
        if run.bold:
            print(run.text)

import os

os.listdir('datasets/')

# Read all the pdf files, extract their contents and store in file data.txt

text =''
for file in os.listdir('datasets/'): 
    if file.endswith('.pdf'): 
        reader = PdfReader('datasets/'+file)
        for page in reader.pages:
            text += page.extract_text()

with open('data.txt','w') as f: 
    f.write(text)




# ======================================================================
# 16 Extractive Text Summarization
# ======================================================================
# #### Install the packages

# [shell/magic - run manually]  pip install gensim sumy

# #### TextRank Algorithm

text = '''The Bengal tiger is a population of the Panthera tigris tigris subspecies. It ranks among the largest of wild cats. It is distributed from India, southern Nepal, Bangladesh, Bhutan to Southwestern China. Its historical range extended to the Indus Basin until the early 19th century, and it is thought to have been present in the Indian subcontinent since the Late Pleistocene about 12,000 to 16,500 years ago. It is threatened by poaching, habitat loss and habitat fragmentation.
As of 2022, the Bengal tiger population was estimated at 3,167–3,682 individuals in India, 316–355 individuals in Nepal, 131 individuals in Bhutan and around 125 individuals in Bangladesh. It is the national animal of India and Bangladesh.
Felis tigris was the scientific name used by Carl Linnaeus in 1758 for the tiger.[1] It was subordinated to the genus Panthera by Reginald Innes Pocock in 1929. Bengal is the traditional type locality of the species and the nominate subspecies Panthera tigris tigris.[2]
The validity of several tiger subspecies in continental Asia was questioned in 1999. Morphologically, tigers from different regions vary little, and gene flow between populations in those regions is considered to have been possible during the Pleistocene. Therefore, it was proposed to recognise only two subspecies as valid, namely P. t. tigris in mainland Asia, and P. t. sondaica in the Greater Sunda Islands and possibly in Sundaland.[3] The nominate subspecies P. t. tigris constitutes two clades: the northern clade comprises the Siberian and Caspian tiger populations, and the southern clade all remaining continental tiger populations.[4] The extinct and living tiger populations in continental Asia have been subsumed to P. t. tigris since the revision of felid taxonomy in 2017.[5]
Results of a genetic analysis of 32 tiger samples indicate that the Bengal tiger samples grouped into a different clade than the Siberian tiger samples.[6]
The Bengal tiger is defined by three distinct mitochondrial nucleotide sites and 12 unique microsatellite alleles. The pattern of genetic variation in the Bengal tiger corresponds to the premise that it arrived in India approximately 12,000 years ago.[7] This is consistent with the lack of tiger fossils from the Indian subcontinent prior to the late Pleistocene, and the absence of tigers from Sri Lanka, which was separated from the subcontinent by rising sea levels in the early Holocene.[8]
A tiger in Bandhavgarh National Park, Madhya Pradesh
The Bengal tiger's coat is yellow to light orange, with stripes ranging from dark brown to black; the belly and the interior parts of the limbs are white, and the tail is orange with black rings. The white tiger is a recessive mutant, which is reported in the wild from time to time in Assam, Bengal, Bihar and especially in the former State of Rewa. However, it is not an occurrence of albinism. In fact, only one fully authenticated case of a true albino tiger is known, and there are no confirmed cases of black tigers, except for a possible specimen examined in Chittagong in 1846.[9] Fourteen Bengal tiger skins in the collection of the Natural History Museum, London have 21–29 stripes.[3] Another recessive mutant is the golden tiger that has a pale golden fur with red-brown stripes.[10] The mutants are very rare in nature.[11]
The greatest skull length of a tiger is 351 mm (13.8 in) in males and 293 mm (11.5 in) in females.[12] It has exceptionally stout teeth. Its canines are 7.5 to 10 cm (3.0 to 3.9 in) long and thus the longest among all cats.[13]
Body weight and size
A male tiger in Jim Corbett National Park, Uttarakhand
A tiger in Kanha National Park, Madhya Pradesh
The Bengal tiger ranks among the biggest wild cats alive.[14] Males and female Bengal tigers in Panna Tiger Reserve reach a head-to-body length of 183–211 cm (72–83 in) and 164–193 cm (65–76 in) respectively, including a tail about 85–110 cm (33–43 in) long. Total length ranges from 283 to 311 cm (111 to 122 in) for male tigers and 255–285 cm (100–112 in) for female tigers.[15] They typically range from 90–110 cm (35–43 in) in shoulder height.[16]
Subadult males weigh between 130 and 170 kg (290 and 370 lb) and reach 200–260 kg (440–570 lb) when adult; subadult females weigh 80–100 kg (180–220 lb) and reach between 110 and 180 kg (240 and 400 lb) when adult.[17] In central India, 42 adult male Bengal tigers weighed on average 190 kg (420 lb) with a range of 167–234 kg (368–516 lb); their total length was 282 cm (111 in) with a range of 267–312 cm (105–123 in), and their average shoulder height was 99 cm (39 in); 39 adult female Bengal tigers weighed an average of 132 kg (291 lb) with a maximum of 156 kg (344 lb) and an average total length of 254 cm (100 in) ranging from 239 to 277 cm (94 to 109 in).[18] Several scientists indicated that adult male Bengal tigers in the Terai consistently attain more than 227 kg (500 lb) of body weight. Seven adult males captured in Chitwan National Park in the early 1970s had an average weight of 235 kg (518 lb) ranging from 200 to 261 kg (441 to 575 lb), and that of the females averaged 140 kg (310 lb) ranging from 116 to 164 kg (256 to 362 lb).[19] Two male tigers captured in Chitwan National Park in the 1980s exceeded weights of 270 kg (600 lb) and are the largest free ranging tigers reported to date.[20]
The smallest recorded weights for Bengal tigers are from the Bangladesh Sundarbans, where adult females weigh 75–80 kg (165–176 lb).[21] Three tigresses from the Bangladesh Sundarbans had a mean weight of 76.7 kg (169 lb). The oldest female weighed 75 kg (165 lb) and was in a relatively poor condition at the time of capture. Their skulls and body weights were distinct from those of tigers in other habitats, indicating that they may have adapted to the unique conditions of the mangrove habitat. Their small sizes are probably due to a combination of intense intraspecific competition and small size of prey available to tigers in the Sundarbans, compared to the larger deer and other prey available to tigers in other parts.[22]
The very large "Leeds Tiger" on display at Leeds City Museum, shot in 1860 near Mussoorie, had a body length of 371 cm (12 ft 2 in) at death. Two tigers shot in Kumaon District and near Oude at the end of the 19th century allegedly measured more than 370 cm (12 ft). But at the time, sportsmen had not yet adopted a standard system of measurement; some measured 'between the pegs' while others measured 'over the curves'.[23] The greatest length of a tiger skull measured 41.3 cm (16.25 in) "over the bone"; this one was shot in the vicinity of Nagina in northern India.[24]
In the beginning of the 20th century, a male tiger was shot in central India with a head and body length of 221 cm (87 in) between pegs, a chest girth of 150 cm (59 in), a shoulder height of 109 cm (43 in) and a tail length of 81 cm (32 in), which was perhaps bitten off by a rival male. This specimen could not be weighed, but it was estimated to weigh about 272 kg (600 lb).[18] A male weighing 259 kg (570 lb) was shot in northern India in the 1930s.[24] A male tiger shot in Nepal weighed 320 kg (710 lb) and measured 328 cm (10 ft 9 in) 'over the curves'.[25] The heaviest wild tiger was possibly a male killed in 1967 in the foothills of the Himalayas that measured 323 cm (127 in) between pegs and 338 cm (133 in) over curves; it weighed 388.7 kg (857 lb) after eating a buffalo calf. Without eating the calf, it would have likely weighed about 324.3 kg (715 lb).[26] In the Central Provinces of India, a male tiger shot weighed 317 kg (699 lb) and measured 3.02 m (9 ft 11 in).[27]
The Bengal tiger rivals the Siberian tiger in average weight.'''

# import the class 
from sumy.summarizers.text_rank import TextRankSummarizer

# import the parser and tokenizer 
from sumy.parsers.plaintext import PlaintextParser
from sumy.nlp.tokenizers import Tokenizer

from nltk.tokenize import sent_tokenize
len(sent_tokenize(text))

# initialize the parser 
parser = PlaintextParser.from_string(text, Tokenizer('english'))

# Find summary in three sentences 
text_rank_summarizer = TextRankSummarizer()

summary = text_rank_summarizer(parser.document, sentences_count=3)

for sent in summary: 
    print(sent,'\n')

# #### Lex Rank Summarizer

from sumy.summarizers.lex_rank import LexRankSummarizer

# Find summary in three sentences 
lex_rank_summarizer = LexRankSummarizer()

lex_summary = lex_rank_summarizer(parser.document, sentences_count=3)

for sent in lex_summary: 
    print(sent,'\n')

# #### LSA Summarizer

from sumy.summarizers.lsa import LsaSummarizer

# Find summary in three sentences 
lsa_summarizer = LsaSummarizer()

lsa_summary = lsa_summarizer(parser.document, sentences_count=3)

for sent in lsa_summary: 
    print(sent,'\n')




# ======================================================================
# 17 Abstractive Text Summarization
# ======================================================================
# [shell/magic - run manually]  pip install transformers

from transformers import pipeline

help(pipeline)

summarizer = pipeline('summarization')

import transformers

transformers.__version__

text = '''The Bengal tiger is a population of the Panthera tigris tigris subspecies. It ranks among the largest of wild cats. It is distributed from India, southern Nepal, Bangladesh, Bhutan to Southwestern China. Its historical range extended to the Indus Basin until the early 19th century, and it is thought to have been present in the Indian subcontinent since the Late Pleistocene about 12,000 to 16,500 years ago. It is threatened by poaching, habitat loss and habitat fragmentation.
As of 2022, the Bengal tiger population was estimated at 3,167–3,682 individuals in India, 316–355 individuals in Nepal, 131 individuals in Bhutan and around 125 individuals in Bangladesh. It is the national animal of India and Bangladesh.
Felis tigris was the scientific name used by Carl Linnaeus in 1758 for the tiger.[1] It was subordinated to the genus Panthera by Reginald Innes Pocock in 1929. Bengal is the traditional type locality of the species and the nominate subspecies Panthera tigris tigris.[2]
The validity of several tiger subspecies in continental Asia was questioned in 1999. Morphologically, tigers from different regions vary little, and gene flow between populations in those regions is considered to have been possible during the Pleistocene. Therefore, it was proposed to recognise only two subspecies as valid, namely P. t. tigris in mainland Asia, and P. t. sondaica in the Greater Sunda Islands and possibly in Sundaland.[3] The nominate subspecies P. t. tigris constitutes two clades: the northern clade comprises the Siberian and Caspian tiger populations, and the southern clade all remaining continental tiger populations.[4] The extinct and living tiger populations in continental Asia have been subsumed to P. t. tigris since the revision of felid taxonomy in 2017.[5]
Results of a genetic analysis of 32 tiger samples indicate that the Bengal tiger samples grouped into a different clade than the Siberian tiger samples.[6]
The Bengal tiger is defined by three distinct mitochondrial nucleotide sites and 12 unique microsatellite alleles. The pattern of genetic variation in the Bengal tiger corresponds to the premise that it arrived in India approximately 12,000 years ago.[7] This is consistent with the lack of tiger fossils from the Indian subcontinent prior to the late Pleistocene, and the absence of tigers from Sri Lanka, which was separated from the subcontinent by rising sea levels in the early Holocene.[8]
A tiger in Bandhavgarh National Park, Madhya Pradesh
The Bengal tiger's coat is yellow to light orange, with stripes ranging from dark brown to black; the belly and the interior parts of the limbs are white, and the tail is orange with black rings. The white tiger is a recessive mutant, which is reported in the wild from time to time in Assam, Bengal, Bihar and especially in the former State of Rewa. However, it is not an occurrence of albinism. In fact, only one fully authenticated case of a true albino tiger is known, and there are no confirmed cases of black tigers, except for a possible specimen examined in Chittagong in 1846.[9] Fourteen Bengal tiger skins in the collection of the Natural History Museum, London have 21–29 stripes.[3] Another recessive mutant is the golden tiger that has a pale golden fur with red-brown stripes.[10] The mutants are very rare in nature.[11]
The greatest skull length of a tiger is 351 mm (13.8 in) in males and 293 mm (11.5 in) in females.[12] It has exceptionally stout teeth. Its canines are 7.5 to 10 cm (3.0 to 3.9 in) long and thus the longest among all cats.[13]
Body weight and size'''

summary = summarizer(text, max_length=60, min_length=20)

print(summary[0]['summary_text'])




# ======================================================================
# 18 Growth Functions
# ======================================================================
import numpy as np 
import matplotlib.pyplot as plt

x = np.arange(-5,5,0.25)

x

y = x * 3.76 + 10.29

y

plt.plot(x, y)

x = np.arange(0,5,0.25)

y = (x **3) * 3.76 + 10.29

plt.plot(x, y)

plt.plot(x, np.log(x))

x = np.arange(-5,5,0.25)

y = 1 / (1 + np.exp(-x))

plt.plot(x, y)




# ======================================================================
# 19 Bag of Words
# ======================================================================
import pandas as pd
from nltk.tokenize import word_tokenize

sent1 = 'It is good practice for us.'
sent2 = 'It was also good to know about it.'

# create the set of distinct words
tokens = word_tokenize(sent1.lower()) + word_tokenize(sent2.lower())
tokens = set(tokens)

# create an empty dataframe
df = pd.DataFrame([], index=[1,2], columns=list(tokens))
df

counts1 = [word_tokenize(sent1.lower()).count(word) for word in df.columns]
counts2 = [word_tokenize(sent2.lower()).count(word) for word in df.columns]

counts2

df.iloc[0,:] = counts1
df.iloc[1,:] = counts2
df

# #### Count Vectorizer

from sklearn.feature_extraction.text import CountVectorizer

cvt = CountVectorizer()

new = cvt.fit_transform([sent1, sent2])

new.toarray()

cvt.get_feature_names_out()

df = pd.DataFrame(new.toarray(), columns=cvt.get_feature_names_out())
df

cvt = CountVectorizer(stop_words='english')

new = cvt.fit_transform([sent1, sent2])
df = pd.DataFrame(new.toarray(), columns=cvt.get_feature_names_out())
df

# #### n-grams

help(CountVectorizer)

cvt = CountVectorizer(ngram_range=(1,2))

new = cvt.fit_transform([sent1, sent2])

cvt.get_feature_names_out()

df = pd.DataFrame(new.toarray(), columns=cvt.get_feature_names_out())
df

cvt = CountVectorizer(ngram_range=(2,2))
new = cvt.fit_transform([sent1, sent2])
df = pd.DataFrame(new.toarray(), columns=cvt.get_feature_names_out())
df

cvt = CountVectorizer(ngram_range=(2,3))
new = cvt.fit_transform([sent1, sent2])
df = pd.DataFrame(new.toarray(), columns=cvt.get_feature_names_out())
df




# ======================================================================
# 20 TF*IDF Vectorizer
# ======================================================================
sent1 = 'It was a beautiful rainy day that made my day awesome.'
sent2 = 'We made it awesome by adding more flavors in it.'

from sklearn.feature_extraction.text import TfidfVectorizer

tfidf = TfidfVectorizer()

new = tfidf.fit_transform([sent1, sent2])

import pandas as pd

df = pd.DataFrame(new.toarray(), columns=tfidf.get_feature_names_out())

df

# File - mreviews.txt
# Download: https://mitu.co.in/dataset
# Read the file, remove the stopwords and create tf*idf vectorize result

f = open('datasets/mreviews.txt')
lines = f.readlines()

tfidf = TfidfVectorizer(stop_words='english')
new = tfidf.fit_transform(lines)
df = pd.DataFrame(new.toarray(), columns=tfidf.get_feature_names_out())

df




# ======================================================================
# 21 Support Vector Machine
# ======================================================================
# #### Import necessary libraries

import pandas as pd 
import numpy as np 
import matplotlib.pyplot as plt 
import seaborn as sns

# #### Load the dataset

df = pd.read_csv('datasets/SMSSpamCollection', sep='\t', 
                 names=['label','text'])

df

# #### Separate the input and output variables

x = df['text']
y = df['label']

# #### Explore the data

set(y)

sns.countplot(x = y);

y.value_counts()

plt.pie(y.value_counts(), 
        labels=y.value_counts().index, autopct='%2.2f%%');

# #### Data Preprocessing

# ##### Data Cleaning

from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import LancasterStemmer, WordNetLemmatizer, PorterStemmer
from nltk import pos_tag

sent = 'Hello friends! How are you? I like Python Programming.'

# tokenize the data 
tokens1 = word_tokenize(sent)

# remove punctuation and numbers 
tokens2 = [token for token in tokens1 if token.isalpha()]

tokens2

# remove stopwords 
tokens3 = [token.lower() for token in tokens2 
           if token.lower() not in stopwords.words('english')]

tokens3

# apply stemming 
ps = PorterStemmer()
tokens4 = [ps.stem(token) for token in tokens3] 
tokens4

def clean_text(sent):
    # tokenize the data 
    tokens1 = word_tokenize(sent)
    # remove punctuation and numbers 
    tokens2 = [token for token in tokens1 if token.isalpha()]
    # remove stopwords 
    tokens3 = [token.lower() for token in tokens2 
           if token.lower() not in stopwords.words('english')]
    tokens4 = [ps.stem(token) for token in tokens3]
    return tokens4

clean_text(sent)

sent1 = 'It was an unexpected rain during the play on ground.'

clean_text(sent1)

x

x.apply(clean_text)

# ##### Vectorization

from sklearn.feature_extraction.text import TfidfVectorizer

tfidf = TfidfVectorizer(analyzer=clean_text)

x_new = tfidf.fit_transform(x)

x_new.shape

tfidf.get_feature_names_out()

len(tfidf.get_feature_names_out())

x_new.toarray()

# #### Cross-Validation

from sklearn.model_selection import train_test_split

x_train, x_test, y_train, y_test = train_test_split(x_new, y, random_state=0)

x_train.shape

x_test.shape

# #### Build the model

from sklearn.svm import SVC

svc_lin = SVC(kernel='linear')

svc_lin.fit(x_train, y_train)

# #### Performance Evaluation

from sklearn.metrics import accuracy_score, classification_report
from sklearn.metrics import ConfusionMatrixDisplay

y_pred = svc_lin.predict(x_test)

ConfusionMatrixDisplay.from_predictions(y_test, y_pred)

print(classification_report(y_test, y_pred))

accuracy_score(y_test, y_pred)

svc_rbf = SVC(kernel='rbf')
svc_rbf.fit(x_train, y_train)
y_pred = svc_rbf.predict(x_test)
accuracy_score(y_test, y_pred)

print(classification_report(y_test, y_pred))

svc_poly = SVC(kernel='poly')
svc_poly.fit(x_train, y_train)
y_pred = svc_poly.predict(x_test)
accuracy_score(y_test, y_pred)

print(classification_report(y_test, y_pred))

svc_sig = SVC(kernel='sigmoid')
svc_sig.fit(x_train, y_train)
y_pred = svc_sig.predict(x_test)
accuracy_score(y_test, y_pred)

print(classification_report(y_test, y_pred))

# #### Prediction on unknown data

# sample.csv
# https://mitu.co.in/dataset 

f = open('datasets/sample.csv')
lines= f.readlines()

lines

new_lines = tfidf.transform(lines)

svc_lin.predict(new_lines)

# #### Save the model and preprocessor

import joblib

joblib.dump(svc_lin,'model.bin')

joblib.dump(tfidf,'preprocess.bin')




# ======================================================================
# 22 Centroid Based Classification
# ======================================================================
# #### Dataset

# UpdatedResumeDataset.csv 
# https://mitu.co.in/dataset

# #### Import necessary libraries

import pandas as pd 
import numpy as np 
import matplotlib.pyplot as plt 
import seaborn as sns

# #### Load the dataset

df = pd.read_csv('datasets/UpdatedResumeDataSet.csv')

df.head()

print(df.iloc[0,1])

# #### Separate the input and output variables

# input
x = df['Resume']

# output 
y = df['Category']

plt.figure(figsize=(16,9))
sns.countplot(y = y);

y.value_counts()

plt.figure(figsize=(10,10))
plt.pie(y.value_counts(), 
        labels=y.value_counts().index, autopct='%2.2f%%');

# #### Define the cleaning function

import re, string

def cleanResume(text): 
    text = re.sub(r'http\S+\s*',' ', text) # URLS
    text = re.sub('RT|cc',' ', text) # RT / CC
    text = re.sub(r'#\S+',' ', text) # Remove hashtage
    text = re.sub(r'@\S+',' ', text) # Remove mentions
    text = re.sub('[%s]' %re.escape(string.punctuation),' ', text) # punctu.. 
    text = re.sub(r'[^\x00-\x7f]',r' ', text) # 
    text = re.sub(r'\s+',' ', text)
    return text

sent = 'हॅलो @Mahesh how are you?\nHave you seen https://fb.com #trend?'

cleanResume(sent)

x[3]

cleanResume(x[3])

cleaned = x.apply(cleanResume)

cleaned

cleaned_data = cleaned.sum()

len(cleaned_data)

# #### Wordcloud visualization

# [shell/magic - run manually]  pip install wordcloud -U

from wordcloud import WordCloud

wcloud = WordCloud(max_words=100).generate(cleaned_data)

plt.figure(figsize=(16,9))
plt.imshow(wcloud)

from nltk.tokenize import word_tokenize

pd.Series(word_tokenize(cleaned_data)).value_counts()

word_tokenize(cleaned_data).count('Exprience Less')

new_data = " ".join(word_tokenize(cleaned_data))

wcloud = WordCloud(max_words=100, collocations=False, 
                   background_color='white').generate(new_data)
plt.figure(figsize=(16,9))
plt.imshow(wcloud)

wcloud.words_

# #### Data Preparation

from sklearn.feature_extraction.text import TfidfVectorizer

tfidf = TfidfVectorizer(stop_words='english')

x_new = tfidf.fit_transform(cleaned)

tfidf.get_feature_names_out()

x_new.shape

# #### Cross Validation

from sklearn.model_selection import train_test_split

x_train, x_test, y_train, y_test = train_test_split(x_new, y, random_state=0)

# #### Build the model

from sklearn.neighbors import NearestCentroid

nc = NearestCentroid()

nc.fit(x_train, y_train)

# #### Performance Evaluation

from sklearn.metrics import classification_report, accuracy_score

y_pred = nc.predict(x_test)

accuracy_score(y_test, y_pred)

print(classification_report(y_test, y_pred))

# #### Predict on unknown data

# Sample_resume.txt
# https://mitu.co.in/dataset

f = open('datasets/Sample_resume.txt')
resume = f.read()

resume = tfidf.transform([resume])

nc.predict(resume)




# ======================================================================
# 23 Text Clustering
# ======================================================================
# headlines.csv
# https://mitu.co.in/dataset

# #### Import necessary libraries

import pandas as pd 
import numpy as np 
import matplotlib.pyplot as plt 
import seaborn as sns

# #### Lead the dataset

df = pd.read_csv('datasets/headlines.csv', names = ['text'])

df.shape

df

print(df['text'][0])

# #### Data Preprocessing

from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk import pos_tag
from nltk.stem import WordNetLemmatizer

swords = stopwords.words('english')

def clean_text(sent): 
    tokens1 = word_tokenize(sent)
    tokens2 = [token for token in tokens1 if token.isalpha()]
    tokens3 = [token.lower() for token in tokens2 
              if token.lower() not in swords]
    return tokens3

sent = 'The raining was expected today. But not completed.'

clean_text(sent)

wnl = WordNetLemmatizer()
wnl.lemmatize('expected', pos='v')

pos_tag(['expected'])

def lemmatize(word): 
    tags = pos_tag([word])
    if tags[0][1].startswith('V'):
        return wnl.lemmatize(word,pos='v')
    if tags[0][1].startswith('N'):
        return wnl.lemmatize(word,pos='n')
    if tags[0][1].startswith('J'):
        return wnl.lemmatize(word,pos='a')
    if tags[0][1].startswith('R'):
        return wnl.lemmatize(word,pos='r')
    return word

lemmatize('smallest')

lemmatize('went')

lemmatize('children')

def clean_text(sent): 
    tokens1 = word_tokenize(sent)
    tokens2 = [token for token in tokens1 if token.isalpha()]
    tokens3 = [lemmatize(token.lower()) for token in tokens2 
              if token.lower() not in swords]
    return tokens3

clean_text(sent)

# #### Vectorize the data

from sklearn.feature_extraction.text import TfidfVectorizer

tfidf = TfidfVectorizer(analyzer=clean_text)

x_new = tfidf.fit_transform(df['text'])

x_new.shape

from wordcloud import WordCloud

final_text = df['text'].apply(clean_text).sum()

final_text = " ".join(final_text)

wcloud = WordCloud().generate(final_text)

plt.figure(figsize=(16,9))
plt.imshow(wcloud)

swords.extend(['say','said','saying','make','made','one','u','new','new','go','gone','went'])

def clean_text(sent): 
    tokens1 = word_tokenize(sent)
    tokens2 = [token for token in tokens1 if token.isalpha()]
    tokens3 = [lemmatize(token.lower()) for token in tokens2 
              if token.lower() not in swords]
    return tokens3

tfidf = TfidfVectorizer(analyzer=clean_text)

x_new = tfidf.fit_transform(df['text'])

x_new.shape

# #### Identify the number of clusters

from sklearn.cluster import KMeans

km = KMeans(n_clusters=10, n_init='auto')

km.fit(x_new)

km.inertia_

sse = []
for k in range(1,16):
    km = KMeans(n_clusters=k, n_init='auto', random_state=0)
    km.fit(x_new)
    sse.append(km.inertia_)

sse

plt.figure(figsize=(16,9))
plt.grid()
plt.xlabel('Value of K')
plt.ylabel('SSE')
plt.plot(range(1,16), sse, marker = 'o', color = 'red')

# [shell/magic - run manually]  pip install kneed

from kneed import KneeLocator

kl = KneeLocator(range(1,16), sse, curve='convex', direction='decreasing')

kl.elbow

# #### Create the clusters

km = KMeans(n_clusters=5, n_init='auto', random_state=0)

labels = km.fit_predict(x_new)

labels

zero = df['text'][labels==0]
one = df['text'][labels==1]
two = df['text'][labels==2]
three = df['text'][labels==3]
four = df['text'][labels==4]

zero.shape, one.shape, two.shape, three.shape, four.shape

zero = " ".join(zero.apply(clean_text).sum())

zero_wcloud = WordCloud(max_words=25).generate(zero)
plt.imshow(zero_wcloud)

zero_wcloud.words_

one = " ".join(one.apply(clean_text).sum())
one_wcloud = WordCloud(max_words=25).generate(one)
plt.imshow(one_wcloud)

two = " ".join(two.apply(clean_text).sum())
two_wcloud = WordCloud(max_words=25).generate(two)
plt.imshow(two_wcloud)

three = " ".join(three.apply(clean_text).sum())
three_wcloud = WordCloud(max_words=25).generate(three)
plt.imshow(three_wcloud)

four = " ".join(four.apply(clean_text).sum())
four_wcloud = WordCloud(max_words=25).generate(four)
plt.imshow(four_wcloud)

new_text = '''PSG will make a move to sign Bournemouth forward Eli Junior Kroupi, who has been linked with Arsenal - L'Equipe
Manchester United fear Marcus Rashford's future might not get resolved until after the World Cup, which would seriously hamper the club's summer transfer business - Daily Mirror
Manchester United and Arsenal target Mateus Fernandes is valued at £80m by relegated West Ham - The Sun
Liverpool transfer target Yan Diomande is favouring a move to Paris Saint-Germain. The young Ivory Coast ace has been identified as an ideal replacement for Mohamed Salah - Foot Mercato
Newcastle have identified Atalanta's Italy right-back Marco Palestra as a summer target but could face competition from Manchester City, Arsenal and Inter Milan - Daily Mail
Also See:
Transfer Centre LIVE!
Download the Sky Sports app
Stream Sky Sports with NOW
Get Sky Sports on WhatsAp
Manchester United are interested in landing Middlesbrough midfielder Hayden Hackney - The Northern Echo
Manchester United have offered Harry Maguire to Inter Milan - La Gazzetta dello Sport
Arsenal teenager Josh Nichols has confirmed he will be leaving the club. The 19-year-old will be departing north London to move to Croatian side NK Kustosija - Daily Mirror
Transfer Centre LIVE! | Latest on YOUR PL club!
Sky Sports Rewards - tickets, offers and more
Download the Sky Sports app for expert analysis, best video & more
Not got Sky? Get Sky Sports or stream with no contract on NOW
Fenerbahce presidential candidate Hakan Safi has ruled out a move for Manchester United target Rafael Leao this summer - Daily Mirror
Tottenham Hotspur have been handed a boost this summer with talented teen Luka Vuskovic focused on the World Cup and not a transfer - talkSPORT
Watch Back Pages on Sky Sports News
Back Pages is a review of the sports headlines from the national newspapers, every Monday to Friday, live on Sky Sports News from 10.30pm.
Missed the show? Catch up on the latest news with the Back Pages podcast.
World Cup
William Saliba could miss the World Cup after flaring up a pre-existing back injury during the Champions League final - The Sun
World football
Pep Guardiola will snub an approach from Sir David Beckham to become the new Inter Miami manager - Daily Mirror
Real Madrid could have to pay up to £13m to secure Jose Mourinho's appointment as manager following the expiration of a break clause in the 63-year-old's contract at Benfica - The Athletic
Saido Berahino has joined the coaching staff of the Burundi national team after retiring from football - The Sun'''

new_text = tfidf.transform([new_text])

km.predict(new_text)




# ======================================================================
# 24 Neural Network Preprocessing
# ======================================================================
lines = [
    'It was a nice rainy day.', 
    'The things were so beautiful on his point.', 
    'When your focus is clear, you won.',
    'Many may happy returns of the day!',
    'One book can change your life.'
]

from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences

tokenizer = Tokenizer()
tokenizer.fit_on_texts(lines)

tokenizer.word_docs

tokenizer.word_counts

tokenizer.word_index

tokenizer.index_word

tokenizer.index_docs

mat = tokenizer.texts_to_matrix(lines)

mat

mat.shape

words_in_first_row = [
    tokenizer.index_word[i] for i, value in enumerate(mat[1]) 
    if value == 1.0 and i != 0
]

words_in_first_row

seq = tokenizer.texts_to_sequences(lines)

seq

padded = pad_sequences(seq, maxlen=10, padding='post')

padded

padded = pad_sequences(seq, maxlen=10, padding='pre')
padded

m = 0
for i in seq: 
    if len(i) > m:
        m = len(i)

m

max([len(i) for i in seq])

padded = pad_sequences(seq, maxlen=max([len(i) for i in seq]), padding='pre')

padded




# ======================================================================
# 25 Artificial Deep Neural Network
# ======================================================================
# #### Import the libraries

import nltk, re, string, os
from nltk.corpus import stopwords
from tensorflow.keras.preprocessing.text import Tokenizer
import numpy as np

# load doc in memory 
def load_doc(filename): 
    f = open(filename)
    text = f.read()
    f.close()
    return text

load_doc('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')

from nltk.tokenize import word_tokenize
swords = stopwords.words('english')

# Function to clean the documents 
def clean_doc(doc):
    tokens = word_tokenize(doc)
    tokens1 = [token for token in tokens if token.isalpha()]
    tokens2 = [token for token in tokens1 if token not in swords]
    tokens3 = [token for token in tokens2 if len(token) > 1]
    return tokens3

d = load_doc('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')
clean_doc(d)

# load the file, clean the data and return the string 
def doc_to_line(filename): 
    doc = load_doc(filename)
    tokens = clean_doc(doc)
    return ' '.join(tokens)

doc_to_line('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')

f = open('datasets/vocab.txt')
vocab = f.read().split()
vocab

# load the file, clean the data and return the string 
def doc_to_line(filename): 
    doc = load_doc(filename)
    tokens = clean_doc(doc)
    tokens = [token for token in tokens if token in vocab]
    return ' '.join(tokens)

doc_to_line('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')

os.listdir('datasets/review_polarity/txt_sentoken/neg')

len(os.listdir('datasets/review_polarity/txt_sentoken/neg'))

len(os.listdir('datasets/review_polarity/txt_sentoken/pos'))

sorted(os.listdir('datasets/review_polarity/txt_sentoken/neg'))

def process_train(directory):
    documents = []
    for filename in os.listdir(directory): 
        if not filename.startswith('cv9'):
            path = directory + '/' + filename
            docs = load_doc(path)
            tokens = clean_doc(docs)
            documents.append(tokens)
    return documents

tr = process_train('datasets/review_polarity/txt_sentoken/neg')

len(tr)

def process_test(directory):
    documents = []
    for filename in os.listdir(directory): 
        if filename.startswith('cv9'):
            path = directory + '/' + filename
            docs = load_doc(path)
            tokens = clean_doc(docs)
            documents.append(tokens)
    return documents

te = process_test('datasets/review_polarity/txt_sentoken/neg')
len(te)

def process_docs(directory, is_train):
    documents = []
    for filename in os.listdir(directory): 
        if is_train and filename.startswith('cv9'):
            continue
        if not is_train and not filename.startswith('cv9'):
            continue
        path = directory + '/' + filename
        docs = load_doc(path)
        tokens = clean_doc(docs)
        documents.append(tokens)
    return documents

tr = process_docs('datasets/review_polarity/txt_sentoken/neg', True)
len(tr)

te = process_docs('datasets/review_polarity/txt_sentoken/neg', False)
len(te)

def load_data(is_train): 
    neg = process_docs('datasets/review_polarity/txt_sentoken/neg', is_train)
    pos = process_docs('datasets/review_polarity/txt_sentoken/pos', is_train) 
    docs = neg + pos 
    labels = [0 for i in range(len(neg))] + [1 for i in range(len(pos))]
    return docs, labels

train_data, train_labels = load_data(True)

len(train_data), len(train_labels)

test_data, test_labels = load_data(False)
len(test_data), len(test_labels)

set(train_labels)

# #### Data Preparation

def create_tokenizer(lines): 
    tokenizer = Tokenizer()
    tokenizer.fit_on_texts(lines) 
    return tokenizer

tokenizer = create_tokenizer(train_data)

len(tokenizer.word_index)

x_train = tokenizer.texts_to_matrix(train_data)
x_test = tokenizer.texts_to_matrix(test_data)

x_train.shape

x_train.shape

# #### Define the model

from keras.models import Sequential
from keras.layers import Input, Dense
from keras.utils import to_categorical, plot_model

# create the model 
model = Sequential()

# input layer
model.add(Input(shape=(36389,)))

# hidden layer
model.add(Dense(256, activation='relu'))

# output layer
model.add(Dense(1, activation='sigmoid'))

# compile the model 
model.compile(loss='binary_crossentropy', optimizer='adam', 
              metrics=['accuracy'])

# #### Train the model

model.fit(x_train, np.array(train_labels), batch_size=5, epochs=10)

plot_model(model, show_dtype=True, show_layer_activations=True, 
           show_layer_names=True, show_shapes=True)

# #### Evaluate on test data

model.evaluate(x_test, np.array(test_labels), batch_size=1)

# #### Prediction of unknown data

review1 = 'Best movie ever! It was great, I will definitely recommend it.'
review2 = 'This is the bad movie. Please dont watch it.'

def predict_sentiment(review):
    tokens = clean_doc(review)
    tokens = ' '.join(tokens)
    encoded = tokenizer.texts_to_matrix([tokens])
    yhat = model.predict(encoded)
    prob = yhat[0,0]
    if prob > 0.5:
        return 'Positive', prob
    return 'Negative', 1-prob

predict_sentiment(review1)

predict_sentiment(review2)




# ======================================================================
# 26 Convolutional Neural Network - Sentiment Analysis
# ======================================================================
# #### Import the libraries

import nltk, re, string, os
from nltk.corpus import stopwords
from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
import numpy as np

# load doc in memory 
def load_doc(filename): 
    f = open(filename)
    text = f.read()
    f.close()
    return text

load_doc('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')

from nltk.tokenize import word_tokenize
swords = stopwords.words('english')

# Function to clean the documents 
def clean_doc(doc):
    tokens = word_tokenize(doc)
    tokens1 = [token for token in tokens if token.isalpha()]
    tokens2 = [token for token in tokens1 if token not in swords]
    tokens3 = [token for token in tokens2 if len(token) > 1]
    tokens3 = [token for token in tokens3 if token in vocab]
    return tokens3

d = load_doc('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')
clean_doc(d)

# load the file, clean the data and return the string 
def doc_to_line(filename): 
    doc = load_doc(filename)
    tokens = clean_doc(doc)
    return ' '.join(tokens)

doc_to_line('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')

f = open('datasets/vocab.txt')
vocab = f.read().split()
vocab

# load the file, clean the data and return the string 
def doc_to_line(filename): 
    doc = load_doc(filename)
    tokens = clean_doc(doc)
    tokens = [token for token in tokens if token in vocab]
    return ' '.join(tokens)

doc_to_line('datasets/review_polarity/txt_sentoken/neg/cv004_12641.txt')

os.listdir('datasets/review_polarity/txt_sentoken/neg')

len(os.listdir('datasets/review_polarity/txt_sentoken/neg'))

len(os.listdir('datasets/review_polarity/txt_sentoken/pos'))

sorted(os.listdir('datasets/review_polarity/txt_sentoken/neg'))

def process_train(directory):
    documents = []
    for filename in os.listdir(directory): 
        if not filename.startswith('cv9'):
            path = directory + '/' + filename
            docs = load_doc(path)
            tokens = clean_doc(docs)
            documents.append(tokens)
    return documents

tr = process_train('datasets/review_polarity/txt_sentoken/neg')

len(tr)

def process_test(directory):
    documents = []
    for filename in os.listdir(directory): 
        if filename.startswith('cv9'):
            path = directory + '/' + filename
            docs = load_doc(path)
            tokens = clean_doc(docs)
            documents.append(tokens)
    return documents

te = process_test('datasets/review_polarity/txt_sentoken/neg')
len(te)

def process_docs(directory, is_train):
    documents = []
    for filename in os.listdir(directory): 
        if is_train and filename.startswith('cv9'):
            continue
        if not is_train and not filename.startswith('cv9'):
            continue
        path = directory + '/' + filename
        docs = load_doc(path)
        tokens = clean_doc(docs)
        tokens = [token for token in tokens if token in vocab]
        documents.append(tokens)
    return documents

tr = process_docs('datasets/review_polarity/txt_sentoken/neg', True)
len(tr)

te = process_docs('datasets/review_polarity/txt_sentoken/neg', False)
len(te)

def load_data(is_train): 
    neg = process_docs('datasets/review_polarity/txt_sentoken/neg', is_train)
    pos = process_docs('datasets/review_polarity/txt_sentoken/pos', is_train) 
    docs = neg + pos 
    labels = [0 for i in range(len(neg))] + [1 for i in range(len(pos))]
    return docs, labels

train_data, train_labels = load_data(True)

len(train_data), len(train_labels)

test_data, test_labels = load_data(False)
len(test_data), len(test_labels)

set(train_labels)

# #### Data Preparation

def create_tokenizer(lines): 
    tokenizer = Tokenizer()
    tokenizer.fit_on_texts(lines) 
    return tokenizer

tokenizer = create_tokenizer(train_data)

len(tokenizer.word_index)

x_train = tokenizer.texts_to_matrix(train_data)
x_test = tokenizer.texts_to_matrix(test_data)

x_train.shape

x_train.shape

def encode_docs(tokenizer, max_length, docs):
    encoded = tokenizer.texts_to_sequences(docs) 
    padded = pad_sequences(encoded, maxlen=max_length, padding='post')
    return padded

max_length = max([len(sent) for sent in train_data])

max_length

x_train = encode_docs(tokenizer, max_length, train_data)
x_test = encode_docs(tokenizer, max_length, test_data)

x_train.shape, x_test.shape

x_train[0]

# #### Define the model

from keras.models import Sequential
from keras.layers import Input, Dense, Conv1D, MaxPool1D, Embedding, Flatten
from keras.utils import to_categorical, plot_model

# create the model 
model = Sequential()

# Embedding layer
model.add(Embedding(len(vocab), 128))

# Convolution layer 
model.add(Conv1D(filters=16, kernel_size=5, activation='relu'))
model.add(MaxPool1D(pool_size=2))

# Flatter layer 
model.add(Flatten())

# hidden layer
model.add(Dense(16, activation='relu'))

# output layer
model.add(Dense(1, activation='sigmoid'))

# compile the model 
model.compile(loss='binary_crossentropy', optimizer='adam', 
              metrics=['accuracy'])

# #### Train the model

model.fit(x_train, np.array(train_labels), batch_size=5, epochs=1)

plot_model(model, show_dtype=True, show_layer_activations=True, 
           show_layer_names=True, show_shapes=True)

# #### Evaluate on test data

model.evaluate(x_test, np.array(test_labels), batch_size=1)

# #### Prediction of unknown data

review1 = 'Best movie ever! It was great, I will definitely recommend it.'
review2 = 'This is the bad movie. Please dont watch it.'

def predict_sentiment(review):
    tokens = clean_doc(review)
    tokens = ' '.join(tokens)
    encoded = encode_docs(tokenizer, max_length, [tokens])
    yhat = model.predict(encoded)
    prob = yhat[0,0]
    if prob > 0.5:
        return 'Positive', prob
    return 'Negative', 1-prob

predict_sentiment(review1)

predict_sentiment(review2)




# ======================================================================
# 27 Sequence Modeling using RNN
# ======================================================================
alphabets = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'

len(alphabets)

dict(enumerate(alphabets))

int_to_char = dict(enumerate(alphabets))

char_to_int = dict((v,i) for i, v in enumerate(alphabets))

char_to_int

# #### Prepare the data

seq_length = 1
x = []
y = []

for i in range(len(alphabets) - seq_length): 
    seq_in = [alphabets[i]] 
    seq_out = alphabets[i+1]
    print(seq_in,'-->',seq_out)
    x += [[char_to_int[seq_in[0]]]]
    y += [char_to_int[seq_out]]

y

x

import numpy as np

x = np.reshape(x, (25,1,1))

x

x_scaled = x / 25

x_scaled

from keras.utils import to_categorical

y_new = to_categorical(y)

y_new

y_new.shape

# #### Build the model

from keras.layers import SimpleRNN, Input, Dense
from keras.models import Sequential

model = Sequential()

model.add(Input((1,1)))
model.add(Dense(32, activation='relu'))
model.add(SimpleRNN(32))
model.add(Dense(26, activation='softmax'))

model.summary()

model.compile(loss='categorical_crossentropy', 
             optimizer='adam', metrics=['accuracy'])

model.fit(x_scaled, y_new, batch_size=1, epochs=500)

model.evaluate(x_scaled, y_new, batch_size=1)

new_char = 'X'
new_char = char_to_int[new_char]
new_char = np.reshape(new_char,(1,1,1))
new_char = new_char / 25
pred = model.predict(new_char, verbose=0)
int_to_char[pred.argmax()]

seq_length = 3
x = []
y = []

for i in range(len(alphabets) - seq_length): 
    seq_in = alphabets[i:i+seq_length] 
    seq_out = alphabets[i+seq_length]
    print(seq_in,'-->',seq_out)
    x.append([char_to_int[char] for char in seq_in])
    y += [char_to_int[seq_out]]

x_new = np.reshape(x, (23, seq_length,1))

y_new = to_categorical(y)

x_new.shape, y_new.shape

x_new = x_new / 25

model = Sequential()

model.add(Input((3,1)))
model.add(Dense(32, activation='relu'))
model.add(SimpleRNN(32))
model.add(Dense(26, activation='softmax'))

model.compile(loss='categorical_crossentropy', 
             optimizer='adam', metrics=['accuracy'])

model.fit(x_new, y_new, batch_size=1, epochs=1000)

new = 'STU'
new_char = np.array([[char_to_int[x]] for x in new])
new_char = np.reshape(new_char,(1,3,1))
new_char = new_char / 25
pred = model.predict(new_char, verbose=0)
int_to_char[pred.argmax()]




# ======================================================================
# 28 LSTM GRU for Classification
# ======================================================================
# [shell/magic - run manually]  pip install tensorflow_datasets

import tensorflow_datasets as tfds
import numpy as np
import tensorflow as tf

imdb, info = tfds.load('imdb_reviews', with_info=True, as_supervised=True)

info

imdb

train_data, test_data = imdb['train'], imdb['test']

training_sentences = []
training_labels = []
testing_sentences = []
testing_labels = []

for s, l in train_data:
    training_sentences.append(str(s.numpy()))
    training_labels.append(l.numpy())

for s, l in test_data:
    testing_sentences.append(str(s.numpy()))
    testing_labels.append(l.numpy())

testing_sentences[0]

testing_labels[0]

len(training_sentences), len(testing_sentences)

len(training_labels), len(testing_labels)

training_labels = np.array(training_labels)
testing_labels  = np.array(testing_labels)

# #### Preprocessing

from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences

help(Tokenizer)

tokenizer = Tokenizer(num_words=10000)

type(training_sentences[0])

tokenizer.fit_on_texts(training_sentences)

word_index = tokenizer.word_index

sequences = tokenizer.texts_to_sequences(training_sentences)
padded = pad_sequences(sequences, maxlen=500,
                       truncating='post', padding='post')
testing_sequences = tokenizer.texts_to_sequences(testing_sentences)
testing_padded = pad_sequences(testing_sequences, maxlen=500,
                       truncating='post', padding='post')

padded.shape

padded[0]

# #### Build the model

from keras.models import Sequential
from keras.layers import Dense, Embedding, SimpleRNN, Bidirectional, LSTM, GRU

# ##### Simple RNN

model_rnn = Sequential([
    Embedding(10000, 50, input_length=500),
    SimpleRNN(32),
    Dense(10, activation='relu'),
    Dense(1, activation='sigmoid')
])

model_rnn.compile(loss='binary_crossentropy', optimizer='adam',
                 metrics=['accuracy'])

training_labels

history = model_rnn.fit(padded, training_labels, epochs=10,
                       validation_data=(testing_padded, testing_labels))

import matplotlib.pyplot as plt

plt.plot(history.history['accuracy'], label = 'accuracy')
plt.plot(history.history['val_accuracy'], label = 'val accuracy')
plt.plot(history.history['loss'], label = 'loss')
plt.title('Model Accuracy')
plt.grid()
plt.legend()

# *LSTM Model*

model_lstm = Sequential([
    Embedding(10000, 50, input_length=500),
    LSTM(32),
    Dense(10, activation='relu'),
    Dense(1, activation='sigmoid')
])

model_lstm.compile(loss='binary_crossentropy', optimizer='adam',
                 metrics=['accuracy'])

history_lstm = model_lstm.fit(padded, training_labels, epochs=10,
                       validation_data=(testing_padded, testing_labels))

plt.plot(history_lstm.history['accuracy'], label = 'accuracy')
plt.plot(history_lstm.history['val_accuracy'], label = 'val accuracy')
plt.plot(history_lstm.history['loss'], label = 'loss')
plt.title('Model Accuracy')
plt.grid()
plt.legend()

# **GRU Model**

model_gru = Sequential([
    Embedding(10000, 50, input_length=500),
    GRU(32),
    Dense(10, activation='relu'),
    Dense(1, activation='sigmoid')
])

model_gru.compile(loss='binary_crossentropy', optimizer='adam',
                 metrics=['accuracy'])

history_gru = model_gru.fit(padded, training_labels, epochs=10,
                       validation_data=(testing_padded, testing_labels))

plt.plot(history_gru.history['accuracy'], label = 'accuracy')
plt.plot(history_gru.history['val_accuracy'], label = 'val accuracy')
plt.plot(history_gru.history['loss'], label = 'loss')
plt.title('Model Accuracy')
plt.grid()
plt.legend()

plt.title('Accuracy Comparison')
plt.plot(history.history['accuracy'], label = 'Simple RNN', color = 'r', marker = '.')
plt.plot(history_lstm.history['accuracy'], label = 'LSTM', color = 'g', marker = '.')
plt.plot(history_gru.history['accuracy'], label = 'GRU', color = 'b', marker = '.')
plt.legend()

new1 = 'This was worst moview. Please do not watch it. WAste of money and time.'
new2 = 'A must watch movie. I lived to watch it. It was just amazing.'




# ======================================================================
# 29 Multi-Label Classification
# ======================================================================
# #### Dataset

# toxic-comments.csv
# https://mitu.co.in/dataset

# #### Import necessary libraries

import pandas as pd 
import numpy as np 
import matplotlib.pyplot as plt 
import seaborn as sns
import re

# #### Load the data

df = pd.read_csv('datasets/toxic-comments.csv')

df.shape

df.columns

df.dtypes

comments = df.sum(numeric_only=True)

comments

comments.sum()

plt.bar(comments.index, comments)

plt.pie(comments, labels=comments.index, autopct='%2.2f%%');

total = df.sum(numeric_only=True, axis=1)

sum(total == 0)

newdf = pd.DataFrame({
    'neutral': [sum(total == 0)],
    'toxic': [sum(total != 0)]
})

newdf.T

# #### Data Cleaning

def clean_text(text): 
    text = text.lower()
    text = re.sub(r'[^a-z0-9]+',' ',text) 
    return text    

sample_text = 'Hello, how are you?'

clean_text(sample_text)

# define input variables 
comments = df['comment_text']

# define output variable
labels = df.drop(['id','comment_text'], axis = 1)

labels

comments_text = comments.apply(clean_text)

comments_text

# #### Prepare the input data

from tensorflow.keras.preprocessing.text import Tokenizer
from tensorflow.keras.preprocessing.sequence import pad_sequences
from keras.models import Sequential
from keras.layers import Dense, Embedding, LSTM, GRU 
from sklearn.model_selection import train_test_split

tokenizer = Tokenizer(num_words=5000)
tokenizer.fit_on_texts(comments_text)
sequences = tokenizer.texts_to_sequences(comments_text)
padded = pad_sequences(sequences, maxlen=200, padding='post')

padded.shape

padded

# #### Cross Validation

x_train, x_test, y_train, y_test = train_test_split(
    padded, labels, random_state=0, test_size=0.2)

x_train.shape, y_train.shape

x_test.shape, y_test.shape

# #### Build the model

model = Sequential()

model.add(Embedding(5000, 100, input_length=200))
model.add(GRU(128))
model.add(Dense(6,activation='sigmoid'))

model.compile(loss='binary_crossentropy', optimizer='adam', 
              metrics=['accuracy'])

model.summary()

model.fit(x_train, y_train, epochs=3, validation_data = (x_test, y_test))




# ======================================================================
# 30 Music Generation
# ======================================================================
import numpy as np 
import tensorflow as tf 
import wave, math, struct
from keras.models import Sequential
from keras.layers import Dense, LSTM, GRU, Activation, Input

# Prepare a dummy data for simulating musical notes 
notes_freqs = {
    'A': 440.0, 'B':493.88, 'C':261.63, 'D':293.66, 'E':393.63, 
    'F': 349.23, 'G': 392.0
}

notes_freqs;

notes = list(notes_freqs.keys())
notes

note_to_int = {note: i for i, note in enumerate(notes)}

note_to_int

int_to_note = {i: note for i, note in enumerate(notes)}
int_to_note

raw_music_data = [notes[np.random.randint(0,7)] for i in range(1000)]

raw_music_data

# #### Data Preparation

seq_length = 3
network_input = []
network_output = []

for i in range(len(raw_music_data) - seq_length): 
    seq_in = raw_music_data[i: i+seq_length]
    seq_out = raw_music_data[i+seq_length]
    network_input.append([note_to_int[char] for char in seq_in])
    network_output.append(note_to_int[seq_out])
    print(seq_in,'-->',seq_out)

n_paterns = len(network_input)

n_paterns

x = np.reshape(network_input,(n_paterns, seq_length, 1))
x

from keras.utils import to_categorical

y = to_categorical(network_output)

y.shape

y

# #### Build the model

model = Sequential()
model.add(Input((3,1)))
model.add(GRU(256))
model.add(Dense(512, activation='relu'))
model.add(Dense(7,activation='softmax'))

model.summary()

model.compile(loss='categorical_crossentropy', 
              optimizer='adam', metrics=['accuracy'])

model.fit(x, y, epochs=1000, batch_size=10)

# #### Generate new melody sequence

start_index = np.random.randint(0, len(network_output))
pattern = network_input[start_index]
pattern

generated_melody = [] 
for i in range(16): 
    x_input = np.reshape(pattern, (1,len(pattern), 1))
    pred = model.predict(x_input, verbose=False)
    index = np.argmax(pred) 
    result = int_to_note[index]
    generated_melody.append(result)
    pattern.append(index)
    pattern = pattern[1:len(pattern)]

generated_melody

# #### Save this as audio file

with wave.open('my_music.wav','w') as wav_file: 
    wav_file.setparams((1,2,44100,0,'NONE','not compressed'))
    for note in generated_melody: 
        freq = notes_freqs[note]
        num_samples = int(0.5 * 44100) 
        for i in range(num_samples): 
            t = float(i) / 44100
            value = int(32767 * 0.5 * math.sin(2*math.pi*freq*t))
            data = struct.pack('<h', value)
            wav_file.writeframes(data)




# ======================================================================
# 31 Word2Vec
# ======================================================================
# [shell/magic - run manually]  pip install gensim -U

# word_corpus.txt, test-glove.txt
# https://mitu.co.in/dataset

# #### Import the libraries

import gensim 
from gensim.models import Word2Vec
import matplotlib.pyplot as plt 
from sklearn.decomposition import PCA

# #### Prapare the data

corpus_text = [
    # --- TOPIC 1: FOOD & FRUITS ---
    "I eat mango and apple everyday",
    "Mango is a very sweet and delicious fruit",
    "Apple pie is a tasty dessert",
    "Banana is yellow and good for health",
    "Fruits like mango apple and banana are healthy",
    "I love eating pizza and burger for dinner",
    "Burger is a fast food",
    "Pizza contains cheese and tomato sauce",
    "I am hungry I want to eat food",
    "Restaurant serves delicious meals",
    "Vegetables and fruits are essential for diet",
    "Cook dinner in the kitchen",
    "Sweet desserts like cake and chocolate",
    
    # --- TOPIC 2: SPORTS & GAMES ---
    "Cricket is a popular sport in India",
    "Virat Kohli plays cricket for the team",
    "Football is played with a ball in a stadium",
    "Messi is a famous football player",
    "Players run fast to score a goal",
    "Tennis is played with a racket and ball",
    "The match was played in a big stadium",
    "The team won the championship trophy",
    "Athletes run in the olympics",
    "Bat and ball are used in cricket matches",
    "Score runs to win the cricket game",
    "The referee blew the whistle",

    # --- TOPIC 3: TECHNOLOGY & CODING ---
    "Python is a popular programming language",
    "Computers use binary code to process data",
    "Artificial Intelligence and machine learning are the future",
    "Data science involves statistics and algorithms",
    "Software developers write code in java and python",
    "The computer requires a keyboard and mouse",
    "Internet connects computers globally",
    "Deep learning models use neural networks",
    "Robots are built using engineering and ai",
    "The server stores massive amounts of data",
    "Programmers debug the software code",
    "Technology is advancing rapidly"
]

# #### Preprocess the data

tokenized_text = []
for sent in corpus_text: 
    tokens = gensim.utils.simple_preprocess(sent)
    tokenized_text.append(tokens)

tokenized_text

model = Word2Vec(sentences=tokenized_text, vector_size=5,
                window=5, min_count=1)

help(Word2Vec)

model.wv['mango']

model.wv['technology']

model.wv.most_similar('fruit')

model.wv.most_similar('mango')

import numpy as np

model.wv.most_similar('fruit')[1]

np.dot(model.wv['banana'], model.wv['fruit']) 

from sklearn.metrics.pairwise import cosine_similarity
cosine_similarity([model.wv['banana']], [model.wv['fruit']])

cosine_similarity([model.wv['computer']], [model.wv['science']])

cosine_similarity([model.wv['python']], [model.wv['language']])

cosine_similarity([model.wv['software']], [model.wv['code']])

cosine_similarity([model.wv['pizza']], [model.wv['burger']])

list1 = ['fruits','banana','football']

model.wv.doesnt_match(list1)

# #### Visualization

words_of_interest = [
    'apple','banana','fruit','burger','pizza', 
    'cricket','football','tennis','player','ball', 
    'python','java','computer','code','software'
]

vectors = [model.wv[word] for word in words_of_interest]

vectors = np.array(vectors)

vectors

vectors.shape

# Reduce the dimensions of vector to 2 
pca = PCA(n_components=2)

vectors_2d = pca.fit_transform(vectors)

vectors_2d

plt.figure(figsize=(16,9))
plt.grid()
plt.ylim(-0.5,0.5)
plt.xlim(-0.5,0.5)
plt.scatter(vectors_2d[:,0], vectors_2d[:,1], color = 'r')
for i, word in enumerate(words_of_interest): 
    plt.annotate(word, xy=(vectors_2d[i,0], vectors_2d[i,1]), 
                ha='right', va='bottom')


# ======================================================================
# 32 GloVe
# ======================================================================
import numpy as np 
from gensim.models import KeyedVectors
from gensim.scripts.glove2word2vec import glove2word2vec

glove2word2vec('datasets/test_glove.txt', 'datasets/test_wv.txt')

model = KeyedVectors.load_word2vec_format('datasets/test_wv.txt', 
                                         binary=False)

model.similarity('king','queen')

model.similarity('mango','apple')

model.similarity('computer','woman')

model.similarity('king','apple')

model.most_similar('keyboard', topn=1)

model.most_similar(positive= ['king','woman'], negative=['man'], topn=1)

# #### Visualize

from sklearn.decomposition import PCA
import matplotlib.pyplot as plt

pca = PCA(n_components=2, random_state=0)

vectors_2d = pca.fit_transform(model.vectors)

vectors_2d

model.index_to_key

plt.figure(figsize=(16,9))
plt.grid()
plt.scatter(vectors_2d[:,0], vectors_2d[:,1], color = 'r')
for i, word in enumerate(model.index_to_key): 
    plt.annotate(word, xy=(vectors_2d[i,0], vectors_2d[i,1]), 
                ha='right', va='bottom')




# ======================================================================
# 33 Pre-Trained Transformers
# ======================================================================
# [shell/magic - run manually]  pip install transformers

from transformers import pipeline

# #### Text Classification

classifier = pipeline('text-classification')

import pandas as pd

text = 'It was possible for him. So he did it.'

result = classifier(text)

df = pd.DataFrame(result)

df

text = 'Nothing is possible without these efforts!'
result = classifier(text)
df = pd.DataFrame(result)
df

text = 'I think it is impossible for him. So just avoid it.'
result = classifier(text)
df = pd.DataFrame(result)
df

text = 'It was so impressive but not upto the mark.'
result = classifier(text)
df = pd.DataFrame(result)
df

text = '''disney's 35th animated feature-- a retooling of the olympian legend crossed with , well , the superman story-- is surprisingly soft at the center . 
great wit , great art , and a great villain ( james woods as hades , lord of the underworld and local lounge act ) can't quite stifle the yawns induced by a bland hero , his colorfully monotonous sidekick ( danny devito as the satyrical trainer phil ) , and a largely unremarkable soundtrack . 
 ( none of the alan menken/david zippel tunes are particular- ly noteworthy . 
some lack lyrical snap . 
others need more memorable melodies . 
boring ballads we expect , but boring production numbers , too ? ) 
so , hercules is a bit of a long sit , but you won't stay bored . 
the highlights include a nifty round of animated action ( herc battling a cgi hydra ) , a steady stream of anachronisms ( " somebody call ix-i-i " ) and pop references ( " let's get ready to rumble ! " ) a la aladdin , and several long-overdue jabs at the mouse's marketing and merchandising depart- ments . 
while not as rock-solid as hunchback , it's a still a new world of improvement over pocohontas . 
directed by ron clements and john musker , with voice credits including tate donovan , susan egan , bob goldthwait , matt frewer , samantha eggar , paul shaffer , and , as lighting bolt zeus , rip torn , who's having a very good summer , also appearing in trial and error and men in black . '''
result = classifier(text)
df = pd.DataFrame(result)
df

text = '''i was going to see ram shrasta on the big screen last night , but before that i stopped into my video store to rent some movies . 
luckily , my video guy was in the middle or recording ram shrasta ! 
i knew it was going to be a bad print with only half-faces and poor sound quality , but i couldn't help myself . . . 
i got it . 
well , after watching it ( it indeed was a bad print ) i was glad i didn't dish out $7 to waste three hours of my time sitting in a bad theatre watching a bad movie . 
this movie really sucks . 
it had so many inconsistencies it was driving me crazy ! ! 
for example , how can jackie shroff grow two feet of hair in just three or four days ? 
jackie shroff looked really stupid with his steven seagal ponytail . 
the songs are bad , the acting is bad ( especially deepti bhatnagar's ) , and the direction is the worst of all . 
the comedy scenes with jagdish and johnny lever just made me want to throw up . 
i didn't even finish the movie . . . 
i couldn't take it anymore . 
note : but if anyone out there liked aatish ( by the same director ) , i guess you'd like this movie . 
it's practically the same except that sanjay dutt is replaced with jackie shroff . 
aditya panscholi is the sidekick again . 
 ( i was going to give it a zero , but i personally like jackie shroff , so i loosened up a little . 
he looked fresh , wore good clothes ( as always ) , and his performance , even though it wasn't his best , was the only saving grace of the movie ) '''
result = classifier(text)
df = pd.DataFrame(result)
df

# #### Named Entity Recognition

ner = pipeline('token-classification')

text = '''Mark Zukerberg will meet Aditya Joshi
on Monday 25 November 2025, 10am for $3 Trillion deal
at Pune for second time.'''

result = ner(text)
df = pd.DataFrame(result)
df

# #### Question-Answering

qa = pipeline('question-answering')

text = 'New Delhi is capital city of India.'
question = 'What is capital city of India?'
result_qa = qa(question=question, context=text)
df = pd.DataFrame([result_qa])
df

text = '''Antarctica (/ænˈtɑːrktɪkə/ ⓘ)[note 1] is Earth's southernmost and least-populated continent. Situated almost entirely south of the Antarctic Circle and surrounded by the Southern Ocean (also known as the Antarctic Ocean), it contains the geographic South Pole. Antarctica is the fifth-largest continent, being about 40% larger than Europe, and has an area of 14,200,000 km2 (5,500,000 sq mi). Most of Antarctica is covered by the Antarctic ice sheet, with an average thickness of 1.9 km (1.2 mi).
Antarctica is, on average, the coldest, driest, and windiest of the continents, and has the highest average elevation. It is mainly a polar desert, with annual precipitation of over 200 mm (8 in) along the coast and far less inland. About 70% of the world's freshwater reserves are frozen in Antarctica, which, if melted, would raise global sea levels by almost 60 metres (200 ft). Antarctica holds the record for the lowest measured temperature on Earth, −89.2 °C (−128.6 °F). The coastal regions can reach temperatures over 10 °C (50 °F) in the summer. Native species of animals include mites, nematodes, penguins, seals and tardigrades. Where vegetation occurs, it is mostly in the form of lichen or moss.'''
question = 'What is lowest temperature on earth?'
result_qa = qa(question=question, context=text)
df = pd.DataFrame([result_qa])
df

# ['audio-classification', 'automatic-speech-recognition', 'conversational', 'depth-estimation', 'document-question-answering', 'feature-extraction', 'fill-mask', 'image-classification', 'image-feature-extraction', 'image-segmentation', 'image-to-image', 'image-to-text', 'mask-generation', 'ner', 'object-detection', 'question-answering', 'sentiment-analysis', 'summarization', 'table-question-answering', 'text-classification', 'text-generation', 'text-to-audio', 'text-to-speech', 'text2text-generation', 'token-classification', 'translation', 'video-classification', 'visual-question-answering', 'vqa', 'zero-shot-audio-classification', 'zero-shot-classification', 'zero-shot-image-classification', 'zero-shot-object-detection', 'translation_XX_to_YY']

# #### Translation

trans = pipeline('translation_en_to_fr')

translator = pipeline(
    "translation", 
    model="facebook/m2m100_418M", 
    src_lang="en", 
    tgt_lang="mr"
)

text = "Hello friends! How are you?"

result = translator(text)

print("Translated Text:", result[0]['translation_text'])

# #### Conversation

from transformers import Conversation

converse = pipeline('conversational')


